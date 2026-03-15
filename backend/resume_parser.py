import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def _parse_pdf(file_bytes: bytes) -> str:
    import pdfplumber
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text.strip())
    full_text = "\n\n".join(text_parts)
    if not full_text.strip():
        raise ValueError("No readable text found in PDF. It may be image-based or password-protected.")
    return full_text


def _parse_docx(file_bytes: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_text = "  |  ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                parts.append(row_text)
    full_text = "\n".join(parts)
    if not full_text.strip():
        raise ValueError("No readable text found in DOCX file.")
    return full_text


def _parse_txt(file_bytes: bytes) -> str:
    try:
        text = file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        text = file_bytes.decode("latin-1")
    if not text.strip():
        raise ValueError("The text file is empty.")
    return text.strip()


def parse_resume(file_bytes: bytes, filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return _parse_pdf(file_bytes)
    elif ext == ".docx":
        return _parse_docx(file_bytes)
    elif ext == ".txt":
        return _parse_txt(file_bytes)
    else:
        raise ValueError(f"Unsupported file type '{ext}'. Upload PDF, DOCX, or TXT.")


def parse_plain_text(text: str) -> str:
    cleaned = text.strip()
    if not cleaned:
        raise ValueError("Pasted text is empty.")
    if len(cleaned) < 50:
        raise ValueError("Text is too short to be a resume. Please paste the full content.")
    return cleaned